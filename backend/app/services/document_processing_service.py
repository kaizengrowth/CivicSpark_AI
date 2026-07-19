import asyncio
import hashlib
import logging
import os
import re
import shutil
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    import magic

    MAGIC_AVAILABLE = True
except ImportError:
    magic = None
    MAGIC_AVAILABLE = False
    print("Warning: python-magic not available, file type detection will be limited")
import pdfplumber

# Document processing imports
import pypdf

try:
    import tiktoken

    TIKTOKEN_AVAILABLE = True
except ImportError:
    tiktoken = None
    TIKTOKEN_AVAILABLE = False
    print("Warning: tiktoken not available, token counting will be limited")
from app.core.config import Settings
from app.core.llm import get_chat_client
from app.models.document import Document, DocumentChunk
from app.models.meeting import AgendaItem
from app.services.agenda_parser import (
    item_chunk_title,
    normalize_item_number,
    parse_agenda_items,
)
from app.services.vector_service import VectorService
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

# Document types that carry agenda-item structure worth preserving.
AGENDA_DOCUMENT_TYPES = {"agenda", "minutes", "meeting_minutes"}


class DocumentProcessor:
    """Handles document text extraction and processing"""

    def __init__(self, settings: Settings):
        self.settings = settings
        self.openai_client, self.chat_model = get_chat_client(settings)
        # tiktoken fetches its vocabulary on first use; fall back to the
        # chars/4 estimate when that isn't possible (offline dev, tests).
        self.encoding = None
        if TIKTOKEN_AVAILABLE:
            try:
                self.encoding = tiktoken.get_encoding("cl100k_base")
            except Exception as e:
                logger.warning(f"tiktoken unavailable, using char estimate: {e}")

    def detect_file_type(self, file_path: str) -> str:
        """Detect file MIME type"""
        if MAGIC_AVAILABLE:
            try:
                return magic.from_file(file_path, mime=True)
            except Exception as e:
                logger.warning(f"Could not detect file type for {file_path}: {e}")

        # Fallback to extension-based detection
        ext = Path(file_path).suffix.lower()
        mime_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
            ".html": "text/html",
            ".htm": "text/html",
        }
        return mime_map.get(ext, "application/octet-stream")

    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using multiple methods for best results"""
        text = ""
        metadata = {"page_count": 0, "extraction_method": ""}

        try:
            # Method 1: Try pdfplumber first (best for structured PDFs)
            with pdfplumber.open(file_path) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)

                if pages:
                    text = "\n\n".join(pages)
                    metadata["page_count"] = len(pdf.pages)
                    metadata["extraction_method"] = "pdfplumber"
                    return text, metadata

        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")

        try:
            # Method 2: Fallback to pypdf
            with open(file_path, "rb") as file:
                pdf_reader = pypdf.PdfReader(file)
                pages = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        pages.append(page_text)

                if pages:
                    text = "\n\n".join(pages)
                    metadata["page_count"] = len(pdf_reader.pages)
                    metadata["extraction_method"] = "pypdf"
                    return text, metadata

        except Exception as e:
            logger.error(f"All PDF extraction methods failed for {file_path}: {e}")

        return text, metadata

    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            paragraphs = [
                paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
            ]
            text = "\n\n".join(paragraphs)

            metadata = {
                "paragraph_count": len(paragraphs),
                "extraction_method": "python-docx",
            }

            return text, metadata

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return "", {}

    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()

            metadata = {
                "line_count": len(text.splitlines()),
                "extraction_method": "plain_text",
            }

            return text, metadata

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="latin-1") as file:
                    text = file.read()

                metadata = {
                    "line_count": len(text.splitlines()),
                    "extraction_method": "plain_text_latin1",
                }

                return text, metadata

            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return "", {}

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return "", {}

    def extract_text(
        self, file_path: str, mime_type: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text based on file type"""
        if mime_type == "application/pdf":
            return self.extract_text_from_pdf(file_path)
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]:
            return self.extract_text_from_docx(file_path)
        elif mime_type == "text/plain":
            return self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {mime_type}")
            return "", {}

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r"\n\s*\n", "\n\n", text)  # Multiple newlines to double newline
        text = re.sub(r"[ \t]+", " ", text)  # Multiple spaces/tabs to single space

        # Remove page headers/footers patterns (common in government docs)
        text = re.sub(r"Page \d+ of \d+", "", text, flags=re.IGNORECASE)
        text = re.sub(r"City of Tulsa.*?\n", "", text, flags=re.IGNORECASE)

        # Remove excessive punctuation
        text = re.sub(r"\.{3,}", "...", text)
        text = re.sub(r"-{3,}", "---", text)

        return text.strip()

    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken"""
        if self.encoding:
            try:
                return len(self.encoding.encode(text))
            except Exception as e:
                logger.warning(f"Error counting tokens: {e}")

        # Fallback: rough estimate (1 token ≈ 4 characters)
        return len(text) // 4

    def chunk_text(
        self, text: str, max_tokens: int = 1000, overlap_tokens: int = 100
    ) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        chunks = []

        # Split by paragraphs first
        paragraphs = text.split("\n\n")

        current_chunk = ""
        current_tokens = 0
        chunk_index = 0

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            paragraph_tokens = self.count_tokens(paragraph)

            # If single paragraph is too long, split it by sentences
            if paragraph_tokens > max_tokens:
                sentences = re.split(r"(?<=[.!?])\s+", paragraph)

                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue

                    sentence_tokens = self.count_tokens(sentence)

                    if current_tokens + sentence_tokens > max_tokens and current_chunk:
                        # Save current chunk
                        chunks.append(
                            {
                                "content": current_chunk.strip(),
                                "chunk_index": chunk_index,
                                "word_count": len(current_chunk.split()),
                                "token_count": current_tokens,
                                "start_char": 0,  # Would need to track this properly
                                "end_char": 0,
                            }
                        )

                        chunk_index += 1

                        # Start new chunk with overlap
                        if overlap_tokens > 0 and current_chunk:
                            overlap_text = self._get_text_overlap(
                                current_chunk, overlap_tokens
                            )
                            current_chunk = overlap_text + " " + sentence
                            current_tokens = self.count_tokens(current_chunk)
                        else:
                            current_chunk = sentence
                            current_tokens = sentence_tokens
                    else:
                        current_chunk += " " + sentence if current_chunk else sentence
                        current_tokens += sentence_tokens
            else:
                # Regular paragraph processing
                if current_tokens + paragraph_tokens > max_tokens and current_chunk:
                    # Save current chunk
                    chunks.append(
                        {
                            "content": current_chunk.strip(),
                            "chunk_index": chunk_index,
                            "word_count": len(current_chunk.split()),
                            "token_count": current_tokens,
                            "start_char": 0,
                            "end_char": 0,
                        }
                    )

                    chunk_index += 1

                    # Start new chunk with overlap
                    if overlap_tokens > 0 and current_chunk:
                        overlap_text = self._get_text_overlap(
                            current_chunk, overlap_tokens
                        )
                        current_chunk = overlap_text + "\n\n" + paragraph
                        current_tokens = self.count_tokens(current_chunk)
                    else:
                        current_chunk = paragraph
                        current_tokens = paragraph_tokens
                else:
                    current_chunk += "\n\n" + paragraph if current_chunk else paragraph
                    current_tokens += paragraph_tokens

        # Add final chunk
        if current_chunk.strip():
            chunks.append(
                {
                    "content": current_chunk.strip(),
                    "chunk_index": chunk_index,
                    "word_count": len(current_chunk.split()),
                    "token_count": current_tokens,
                    "start_char": 0,
                    "end_char": 0,
                }
            )

        return chunks

    def _get_text_overlap(self, text: str, overlap_tokens: int) -> str:
        """Get the last N tokens worth of text for overlap"""
        words = text.split()
        # Rough estimate: take last N/4 words for overlap
        overlap_words = max(1, overlap_tokens // 4)
        return " ".join(words[-overlap_words:])

    async def extract_keywords(self, text: str) -> List[str]:
        """Extract keywords using OpenAI"""
        if not self.openai_client:
            return []

        try:
            # Truncate text if too long
            max_chars = 3000
            if len(text) > max_chars:
                text = text[:max_chars] + "..."

            response = self.openai_client.chat.completions.create(
                model=self.chat_model,
                messages=[
                    {
                        "role": "system",
                        "content": "Extract 5-10 key terms and phrases from this Tulsa city document. Return as a comma-separated list.",
                    },
                    {"role": "user", "content": f"Document excerpt:\n\n{text}"},
                ],
                max_tokens=100,
                temperature=0.3,
            )

            keywords_text = response.choices[0].message.content.strip()
            keywords = [kw.strip() for kw in keywords_text.split(",") if kw.strip()]

            return keywords[:10]  # Limit to 10 keywords

        except Exception as e:
            logger.error(f"Error extracting keywords: {e}")
            return []

    async def generate_summary(self, text: str) -> str:
        """Generate document summary using OpenAI"""
        if not self.openai_client:
            return ""

        try:
            # Truncate text if too long
            max_chars = 4000
            if len(text) > max_chars:
                text = text[:max_chars] + "..."

            response = self.openai_client.chat.completions.create(
                model=self.chat_model,
                messages=[
                    {
                        "role": "system",
                        "content": "Summarize this Tulsa city document in 2-3 sentences. Focus on key decisions, policies, or information relevant to Tulsa citizens.",
                    },
                    {"role": "user", "content": f"Document:\n\n{text}"},
                ],
                max_tokens=150,
                temperature=0.3,
            )

            return response.choices[0].message.content.strip()

        except Exception as e:
            logger.error(f"Error generating summary: {e}")
            return ""


def _sha256_file(file_path: str) -> str:
    """Content hash of a source file, for provenance and dedup"""
    digest = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(1 << 20), b""):
            digest.update(block)
    return digest.hexdigest()


class DocumentProcessingService:
    """Main service for processing documents and adding them to RAG system"""

    def __init__(self, settings: Settings, db: Session):
        self.settings = settings
        self.db = db
        self.processor = DocumentProcessor(settings)
        self.vector_service = VectorService(settings, db)

    async def process_document(
        self, file_path: str, document_data: Dict[str, Any]
    ) -> Optional[Document]:
        """Process a document file and add it to the database and vector store"""
        try:
            # Provenance first: hash the source so re-uploads of unchanged
            # files are deduplicated and every document records what was
            # ingested and when.
            content_hash = _sha256_file(file_path)
            existing = (
                self.db.query(Document)
                .filter(Document.content_hash == content_hash)
                .first()
            )
            if existing is not None:
                logger.info(
                    f"Skipping ingest: identical content already indexed as "
                    f"document {existing.id}"
                )
                return existing

            # Detect file type
            mime_type = self.processor.detect_file_type(file_path)

            # Extract text
            text, extraction_metadata = self.processor.extract_text(
                file_path, mime_type
            )

            if not text.strip():
                logger.error(f"No text extracted from {file_path}")
                return None

            # Clean text
            cleaned_text = self.processor.clean_text(text)

            # Generate AI analysis
            summary = await self.processor.generate_summary(cleaned_text[:4000])
            keywords = await self.processor.extract_keywords(cleaned_text[:3000])

            # Persist the original file: callers pass a temp path that is
            # deleted after upload, so keep our own copy for reprocessing.
            stored_path = self._store_file(file_path)

            # Create document record
            meeting_id = document_data.get("meeting_id")
            document = Document(
                title=document_data.get("title", Path(file_path).stem),
                meeting_id=meeting_id,
                content=cleaned_text,
                summary=summary,
                document_type=document_data.get("document_type", "unknown"),
                category=document_data.get("category", ""),
                source_url=document_data.get("source_url", ""),
                content_hash=content_hash,
                retrieved_at=datetime.now(timezone.utc),
                file_path=stored_path,
                file_name=Path(file_path).name,
                file_size=os.path.getsize(file_path),
                mime_type=mime_type,
                document_date=document_data.get("document_date"),
                effective_date=document_data.get("effective_date"),
                word_count=len(cleaned_text.split()),
                page_count=extraction_metadata.get("page_count", 0),
                keywords=keywords,
                tags=document_data.get("tags", []),
                is_public=document_data.get("is_public", True),
                uploaded_by=document_data.get("uploaded_by"),
                processing_status="processing",
            )

            # Save to database
            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)

            # Chunk the text: by agenda item when the document has that
            # structure, fixed windows otherwise.
            chunks_data = self._chunk_document(cleaned_text, document.document_type)

            # Match parsed item numbers to structured AgendaItem rows so
            # chunks carry real identity keys, not just labels.
            item_id_by_number = {}
            if meeting_id:
                agenda_rows = (
                    self.db.query(AgendaItem)
                    .filter(AgendaItem.meeting_id == meeting_id)
                    .all()
                )
                item_id_by_number = {
                    normalize_item_number(row.item_number): row.id
                    for row in agenda_rows
                    if row.item_number
                }

            # Create chunk records and prepare for vector store
            chunks_for_vector = []
            for chunk_data in chunks_data:
                item_number = chunk_data.get("item_number")
                chunk = DocumentChunk(
                    document_id=document.id,
                    content=chunk_data["content"],
                    chunk_index=chunk_data["chunk_index"],
                    word_count=chunk_data["word_count"],
                    start_char=chunk_data["start_char"],
                    end_char=chunk_data["end_char"],
                    section_title=chunk_data.get("section_title"),
                    section_type=chunk_data.get("section_type"),
                    item_number=item_number,
                    meeting_id=meeting_id,
                    agenda_item_id=(
                        item_id_by_number.get(item_number) if item_number else None
                    ),
                )

                self.db.add(chunk)

                # Prepare for vector store
                chunks_for_vector.append(
                    {
                        "document_id": document.id,
                        "chunk_index": chunk_data["chunk_index"],
                        "content": chunk_data["content"],
                        "document_type": document.document_type,
                        "category": document.category,
                        "word_count": chunk_data["word_count"],
                    }
                )

            # Update document with chunk count
            document.chunk_count = len(chunks_data)
            document.processing_status = "completed"
            document.is_processed = True

            self.db.commit()

            # Add to vector store
            success = await self.vector_service.add_document_chunks(chunks_for_vector)

            if not success:
                logger.error(f"Failed to add document {document.id} to vector store")
                document.processing_status = "failed"
                document.processing_error = "Failed to add to vector store"
                self.db.commit()

            return document

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")

            # Update document status if it exists
            if "document" in locals():
                document.processing_status = "failed"
                document.processing_error = str(e)
                self.db.commit()

            return None

    def _chunk_document(self, cleaned_text: str, document_type: str) -> list:
        """Chunk by agenda item when structure is present, else fixed windows"""
        if (document_type or "").lower() in AGENDA_DOCUMENT_TYPES:
            items = parse_agenda_items(cleaned_text)
            if items:
                logger.info(f"Parsed {len(items)} agenda items; chunking by item")
                return self._chunk_by_items(items)
            logger.info("No agenda structure found; falling back to fixed windows")
        return self.processor.chunk_text(cleaned_text)

    def _chunk_by_items(self, items: list) -> list:
        """One or more chunks per agenda item, item metadata on every chunk"""
        chunks = []
        for item in items:
            body = f"{item.title}\n\n{item.text}".strip()
            sub_chunks = self.processor.chunk_text(body) or [
                {
                    "content": body,
                    "chunk_index": 0,
                    "word_count": len(body.split()),
                    "start_char": 0,
                    "end_char": 0,
                }
            ]
            for sub in sub_chunks:
                sub["chunk_index"] = len(chunks)
                sub["section_title"] = item_chunk_title(item)
                sub["section_type"] = "agenda_item"
                sub["item_number"] = item.item_number
                chunks.append(sub)
        return chunks

    def _store_file(self, file_path: str) -> str:
        """Copy an uploaded file into the configured storage directory"""
        try:
            storage_dir = Path(self.settings.storage_dir) / "documents"
            storage_dir.mkdir(parents=True, exist_ok=True)
            suffix = Path(file_path).suffix
            destination = storage_dir / f"{uuid.uuid4().hex}{suffix}"
            shutil.copy2(file_path, destination)
            return str(destination)
        except Exception as e:
            logger.warning(f"Could not store original file {file_path}: {e}")
            return file_path

    async def reprocess_document(self, document_id: int) -> bool:
        """Reprocess an existing document"""
        try:
            document = (
                self.db.query(Document).filter(Document.id == document_id).first()
            )
            if not document:
                return False

            # Delete existing chunks
            self.db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).delete()

            # Delete from vector store
            await self.vector_service.delete_document_chunks(document_id)

            # Reprocess
            if document.file_path and os.path.exists(document.file_path):
                document_data = {
                    "title": document.title,
                    "document_type": document.document_type,
                    "category": document.category,
                    "source_url": document.source_url,
                    "document_date": document.document_date,
                    "effective_date": document.effective_date,
                    "tags": document.tags,
                    "is_public": document.is_public,
                    "uploaded_by": document.uploaded_by,
                }

                processed_doc = await self.process_document(
                    document.file_path, document_data
                )
                return processed_doc is not None

            return False

        except Exception as e:
            logger.error(f"Error reprocessing document {document_id}: {e}")
            return False
